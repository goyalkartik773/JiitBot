"""
JIIT Project Synopsis Generator
================================

Automated tool for generating professional project synopses and presentations
in JIIT (Jaypee Institute of Information Technology) format.

Features:
---------
- AI-powered content generation using Google Gemini
- Generates both PDF and DOCX formats
- Professional JIIT-branded cover page
- Automatic sections: Objective, Introduction, Key Features, Technologies, References
- Support for custom project images and diagrams
- Proper text wrapping and pagination

Components:
-----------
- draw_section(): Helper function for PDF content rendering with proper wrapping
- call_gemini_ai(): Interfaces with Google Gemini API for content generation
- create_project_pdf(): Generates PDF document with ReportLab
- create_project_doc(): Generates DOCX document with python-docx
- show(): Main Streamlit UI function

Author: Kartik, Manav, Sujal
Supervisor: Dr. Tribhuvan Kumar Tewary
"""

# Streamlit for web interface
import streamlit as st

# Standard library imports
import io
import json
import re
import os

# Google Generative AI for content generation
import google.generativeai as genai

# ReportLab for PDF generation
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib.utils import ImageReader
from reportlab.platypus import Paragraph, ListFlowable, ListItem
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT

# PIL for image processing
from PIL import Image

# Package utilities
import pkg_resources

# python-docx for DOCX generation
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ============================================================================
# PDF HELPER FUNCTIONS
# ============================================================================

def draw_section(c, y_pos, width, content_data, title):
    """
    Draws a complete section in the PDF with proper text wrapping and pagination.
    
    This function handles complex text rendering including:
    - Automatic text wrapping for long content
    - Page breaks when content doesn't fit
    - Support for strings, lists, and nested dictionaries
    - Proper spacing and indentation
    
    Args:
        c: ReportLab Canvas object for drawing
        y_pos (float): Top Y-coordinate where section starts (in points)
        width (float): Page width for calculating content width
        content_data: Content to render (str, list, or dict)
        title (str): Section title to display
    
    Returns:
        float: New y_pos (top coordinate for next section)
    
    Implementation Notes:
        - Uses ReportLab's Paragraph and ListFlowable for proper wrapping
        - Automatically creates new pages when content doesn't fit
        - Maintains consistent margins and spacing throughout
    """
    # Get standard styles
    styles = getSampleStyleSheet()
    
    # --- Define Margins and Page Geometry ---
    margin_x = inch
    page_bottom_margin = inch
    content_width = width - (2 * margin_x)
    
    # --- 1. Draw Title ---
    styleH = styles['h2'] # Use 'h2' or 'h1'
    p_title = Paragraph(title, styleH)
    
    # Calculate height
    # 'avail_height' is the space left on the page
    avail_height = y_pos - page_bottom_margin
    w_title, h_title = p_title.wrapOn(c, content_width, avail_height)
    
    # Check for page break (if title alone doesn't fit)
    if h_title > avail_height:
        c.showPage()
        y_pos = letter[1] - margin_x # Reset to top of new page
        avail_height = y_pos - page_bottom_margin
        w_title, h_title = p_title.wrapOn(c, content_width, avail_height)

    # Draw the title (y is the BOTTOM of the flowable)
    draw_y = y_pos - h_title
    p_title.drawOn(c, margin_x, draw_y)
    
    # Update our TOP cursor
    y_pos = draw_y - 0.1 * inch # Move down, adding a small gap

    # --- 2. Draw Content ---
    content_flowables = []
    styleN = styles['Normal']
    # Define a style for list items with minimal indentation and reduced spacing
    styleL = ParagraphStyle(
        name='list', 
        parent=styleN, 
        leftIndent=0,  # No indent on paragraph itself
        spaceBefore=0,
        spaceAfter=1,
        leading=11  # Tight line spacing
    )

    if isinstance(content_data, str):
        p_text = Paragraph(content_data.replace('\n', '<br/>'), styleN)
        content_flowables.append(p_text)
        
    elif isinstance(content_data, list):
        list_items = []
        for item_text in content_data:
            # Create a Paragraph for each list item to allow wrapping
            p_item = Paragraph(item_text, styleL)
            list_items.append(p_item)
        # Use ListFlowable with minimal indentation - bullet closer to text
        list_flow = ListFlowable(
            list_items, 
            bulletType='bullet', 
            start='bulletchar',
            bulletOffsetY=-1,
            leftIndent=12,  # Reduced from 18 - bullet closer to margin
            bulletIndent=8,  # Bullet position - closer to text
            rightIndent=0
        )
        content_flowables.append(list_flow)
        
    elif isinstance(content_data, dict):
        for category, items in content_data.items():
            # Add the category as a bolded paragraph
            p_category = Paragraph(f"<b>{category}</b>", styleN)
            content_flowables.append(p_category)
            
            list_items_for_dict = []
            if isinstance(items, list):
                for item_text in items:
                    p_item = Paragraph(item_text, styleL)
                    list_items_for_dict.append(p_item)
            else:
                # Ensure even non-list items are wrapped
                p_item = Paragraph(str(items), styleL)
                list_items_for_dict.append(p_item)
            
            # Create a ListFlowable with minimal indentation
            item_list_flow = ListFlowable(
                list_items_for_dict, 
                bulletType='bullet', 
                start='bulletchar',
                bulletOffsetY=-1,
                leftIndent=12,  # Reduced indentation
                bulletIndent=8,  # Bullet closer to text
                rightIndent=0
            )
            content_flowables.append(item_list_flow)

    # Now draw all the content flowables, checking for page breaks
    for item in content_flowables:
        avail_height = y_pos - page_bottom_margin
        w, h = item.wrapOn(c, content_width, avail_height) # Calculate height
        
        # Check if it fits
        if h > avail_height:
            c.showPage()
            y_pos = letter[1] - margin_x # Reset to top
            avail_height = y_pos - page_bottom_margin
            w, h = item.wrapOn(c, content_width, avail_height)
        
        # Draw it (y is the BOTTOM)
        draw_y = y_pos - h
        item.drawOn(c, margin_x, draw_y)
        
        # Update our TOP cursor with reduced spacing
        y_pos = draw_y - 0.05 * inch # Smaller gap between items

    y_pos -= 0.1 * inch  # Reduced extra space after the whole section
    return y_pos

    
def call_gemini_ai(title):
    """
    Calls Google Gemini API to generate project synopsis content.
    
    This function:
    1. Configures the Gemini API with the provided key
    2. Constructs a detailed prompt for generating synopsis sections
    3. Parses and validates the JSON response
    4. Returns structured content for the synopsis
    
    Args:
        title (str): Project title to generate content for
    
    Returns:
        dict: Generated content with keys:
            - objective: Project objective (2-3 sentences)
            - introduction: Detailed introduction (2-3 paragraphs)
            - key_features: List of 5-7 key features
            - technologies_used: Dict of technology categories and descriptions
            - references: Dict of reference categories and lists
        None: If API call fails or response is invalid
    
    Note:
        - Uses gemini-flash-latest model for fast generation
        - Implements error handling for API failures
        - Validates response structure before returning
    """
    # Configure the Gemini API with the API key from environment variables
    # SECURITY: Never hardcode API keys in your code!
    try:
        # Try to get API key from Streamlit secrets first
        if hasattr(st, 'secrets') and 'GEMINI_API_KEY' in st.secrets:
            api_key = st.secrets['GEMINI_API_KEY']
        # Fallback to environment variable
        elif os.getenv('GEMINI_API_KEY'):
            api_key = os.getenv('GEMINI_API_KEY')
        else:
            st.error("""
            ⚠️ **Gemini API Key Not Found!**
            
            Please add your API key in one of these ways:
            
            **Option 1: Streamlit Secrets (Recommended)**
            1. Create `.streamlit/secrets.toml` file
            2. Add: `GEMINI_API_KEY = "your-api-key-here"`
            
            **Option 2: Environment Variable**
            1. Create `.env` file
            2. Add: `GEMINI_API_KEY=your-api-key-here`
            
            Get your API key from: https://makersuite.google.com/app/apikey
            """)
            return None
        
        genai.configure(api_key=api_key)
    except Exception as e:
        st.error(f"API Key configuration error: {e}")
        return None

    # Use the Gemini Flash model for fast content generation
    model_name = "gemini-flash-latest"
    
    prompt = f"""
    You are an expert academic writer. A student is making a project synopsis for their college.
    The project title is: "{title}"
    
    Generate the following sections for the synopsis:
    - objective: A 2-3 sentence paragraph.
    - introduction: A detailed 2-3 paragraph introduction.
    - key_features: A Python list of 5-7 string features.
    - technologies_used: A Python dictionary where keys are categories (e.g., "Web Framework", "Database") and values are string descriptions.
    - references: A Python dictionary where keys are categories (e.g., "Books", "Web Resources") and values are Python lists of string references.

    You MUST return ONLY a single, valid JSON object with no other text or markdown formatting.
    The JSON keys must be: "objective", "introduction", "key_features", "technologies_used", "references".

    Example format:
    {{
        "objective": "The primary objective of this project is to...",
        "introduction": "The Student Management System is a comprehensive platform...",
        "key_features": ["Secure User Authentication", "Dashboard for Admins", "Automated Attendance Tracking"],
        "technologies_used": {{ 
            "Language": "Python", 
            "Web Framework": "Streamlit",
            "Database": "MySQL" 
        }},
        "references": {{ 
            "Web Resources": ["https://docs.streamlit.io", "https://www.python.org"] 
        }}
    }}
    """
    
    try:
        model = genai.GenerativeModel(model_name)
        response = model.generate_content(prompt)
        
        # Clean the response to find the JSON
        json_match = re.search(r'\{.*\}', response.text, re.DOTALL)
        if not json_match:
            st.error("AI Response Error: No valid JSON object found.")
            st.text_area("AI Raw Response (for debugging)", response.text, height=150)
            return None
            
        json_text = json_match.group(0)
        ai_content = json.loads(json_text)
        
        # Validate required keys
        required_keys = ["objective", "introduction", "key_features", "technologies_used", "references"]
        if not all(key in ai_content for key in required_keys):
            st.error("AI Response Error: The generated JSON is missing one or more required keys.")
            st.json(ai_content)
            return None

        return ai_content
        
    except Exception as e:
        st.error(f"Error calling Gemini API with model '{model_name}': {e}")
        if 'response' in locals():
            st.text_area("AI Raw Response (for debugging)", response.text, height=150)
        return None

def create_project_pdf(user_data, ai_data, uploaded_images):
    """
    Creates a complete PDF document for the project synopsis.
    
    The PDF includes:
    1. Cover page with JIIT branding, logo, and project details
    2. Team member information in tabular format
    3. Guide/supervisor information
    4. AI-generated content sections (objective, introduction, etc.)
    5. Optional image pages with headings
    
    Args:
        user_data (dict): User-provided information containing:
            - title: Project title
            - category: Project category/lab
            - team: List of dicts with 'name' and 'enroll' keys
            - guides: List of guide names
        ai_data (dict): AI-generated content with sections
        uploaded_images (list): List of image file objects to include
    
    Returns:
        io.BytesIO: PDF file buffer ready for download
    
    Implementation:
        - Uses ReportLab for PDF generation
        - Implements automatic pagination
        - Handles image scaling and centering
        - Maintains JIIT branding standards
    """
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter # (8.5 * 72, 11 * 72)
    
    # === Page 1: Cover Page (Accurate Layout) ===
    
    c.setFont("Helvetica-Bold", 16)
    c.drawCentredString(width / 2.0, 10.5 * inch, "JAYPEE INSTITUTE OF INFORMATION TECHNOLOGY")
    c.setFont("Helvetica-Bold", 14)
    c.drawCentredString(width / 2.0, 10.3 * inch, "NOIDA, SEC-62")
    
    # --- NEW: Project Category/Lab from user input ---
    c.setFont("Helvetica-Bold", 14)
    # Positioning this relative to the address or the logo, adjust as needed
    c.drawCentredString(width / 2.0, 9.8 * inch, user_data['category'].upper()) 
    # --- END NEW ---
    
    # Draw JIIT Logo
    try:
        # Assumes 'jiit_logo.png' is in the same folder
        c.drawImage("jiit_logo.png", width / 2.0 - inch, 6.2 * inch, 
                    width=2 * inch, preserveAspectRatio=True)
    except Exception as e:
        c.drawCentredString(width / 2.0, 9 * inch, "[Logo 'jiit_logo.png' not found]")
        
    c.setFont("Helvetica-Bold", 16)
    c.drawCentredString(width / 2.0, 5.65 * inch, "Project Synopsis")
    
    c.setFont("Helvetica-Bold", 14)
    c.drawCentredString(width / 2.0, 5.3 * inch, f"TITLE: {user_data['title']}")

    c.setFont("Helvetica-Bold", 14)
    c.drawCentredString(width / 2.0, 5.0 * inch, "SUBMITTED BY:")
    
    # --- Accurate Team Member Drawing ---
    c.setFont("Helvetica-Bold", 12)
    # Set X-coordinates for the columns
    name_x = 2 * inch
    enroll_x = 5.5 * inch
    
    # Draw Headers
    c.drawString(name_x, 4.5 * inch, "NAME")
    c.drawString(enroll_x, 4.5 * inch, "ENROLLMENT")
    # Draw line under headers
    c.line(name_x - 0.1 * inch, 4.4 * inch, enroll_x + 1.5 * inch, 4.4 * inch)
    
    c.setFont("Helvetica", 12)
    y_team = 4.1 * inch # Starting Y for the first member
    for member in user_data['team']:
        c.drawString(name_x, y_team, member['name'])
        c.drawString(enroll_x, y_team, member['enroll'])
        y_team -= 0.3 * inch # Move down for the next member
    # --- End Team Section ---

    # Draw Guides at the bottom
    c.setFont("Helvetica-Bold", 14)
    c.drawString(1.5 * inch, 3.0 * inch, "Submitted to:")
    c.setFont("Helvetica", 12)
    y_guides = 2.7 * inch
    for guide in user_data['guides']:
        c.drawString(1.5 * inch, y_guides, guide.strip())
        y_guides -= 0.3 * inch
        
    c.showPage()
    
    # === AI Content Pages ===
  #  y_pos = 10 * inch # Start at top of new page
    y_pos = letter[1] - inch
    
    y_pos = draw_section(c, y_pos, width, ai_data['objective'], "OBJECTIVE:")
    y_pos = draw_section(c, y_pos, width, ai_data['introduction'], "INTRODUCTION:")
    y_pos = draw_section(c, y_pos, width, ai_data['key_features'], "KEY FEATURES:")
    y_pos = draw_section(c, y_pos, width, ai_data['technologies_used'], "TECHNOLOGIES USED:")
    
    # References might need its own page
    if y_pos < 5 * inch:
        c.showPage()
        y_pos = 10 * inch
    y_pos = draw_section(c, y_pos, width, ai_data['references'], "REFERENCES:")

    # === New Image Pages (Optional) ===
    if uploaded_images:
        for img_file in uploaded_images:
            c.showPage() # Create a new, separate page
            
            # 1. Draw the "Relevant Heading" using the filename
            c.setFont("Helvetica-Bold", 16)
            heading = img_file.name.split('.')[0].replace('_', ' ').title()
            c.drawCentredString(width / 2.0, 10.5 * inch, heading)
            
            # 2. Draw the image
            try:
                # Reset file pointer and read bytes
                img_file.seek(0)
                img_bytes = img_file.read()
                
                # Open image with PIL to get dimensions
                img_io = io.BytesIO(img_bytes)
                img = Image.open(img_io)
                img_width, img_height = img.size
                
                # Scale image to fit page (max 7.5" wide, 9" tall)
                max_w = 7.5 * inch
                max_h = 9 * inch
                ratio = min(max_w / img_width, max_h / img_height)
                
                new_w = img_width * ratio
                new_h = img_height * ratio
                
                # Center the image
                x = (width - new_w) / 2.0
                y = (height - new_h) / 2.0 - (0.5 * inch) # Center vertically with space for title
                
                # Save PIL Image to BytesIO in PNG format for reportlab
                img_buffer = io.BytesIO()
                # Convert to RGB if necessary (for JPEG images)
                if img.mode in ('RGBA', 'LA', 'P'):
                    rgb_img = Image.new('RGB', img.size, (255, 255, 255))
                    if img.mode == 'P':
                        img = img.convert('RGBA')
                    rgb_img.paste(img, mask=img.split()[-1] if img.mode in ('RGBA', 'LA') else None)
                    img = rgb_img
                elif img.mode != 'RGB':
                    img = img.convert('RGB')
                
                img.save(img_buffer, format='PNG')
                img_buffer.seek(0)
                
                # Use ImageReader to wrap BytesIO for reportlab
                img_reader = ImageReader(img_buffer)
                c.drawImage(img_reader, x, y, width=new_w, height=new_h)
                
            except Exception as e:
                c.drawCentredString(width/2.0, height/2.0, f"[Error loading image: {e}]")

    # --- Finalize PDF ---
    c.save()
    buffer.seek(0)
    return buffer

def create_project_doc(user_data, ai_data, uploaded_images):
    """
    Creates the complete DOCX document from user and AI data.
    Returns a BytesIO buffer.
    """
    document = docx.Document()
    
    # === Page 1: Cover Page ===
    
    # Title: JAYPEE INSTITUTE...
    title_para = document.add_paragraph("JAYPEE INSTITUTE OF INFORMATION TECHNOLOGY")
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.runs[0] if title_para.runs else title_para.add_run()
    title_run.bold = True
    title_run.font.size = Pt(16)
    
    # Address
    address_para = document.add_paragraph("NOIDA, SEC-62")
    address_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    address_run = address_para.runs[0] if address_para.runs else address_para.add_run()
    address_run.bold = True
    address_run.font.size = Pt(14)
    
    # Category
    category_para = document.add_paragraph(user_data['category'].upper())
    category_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    category_run = category_para.runs[0] if category_para.runs else category_para.add_run()
    category_run.bold = True
    category_run.font.size = Pt(14)
    
    # Add spacing
    document.add_paragraph()
    
    # Add JIIT Logo
    try:
        logo_para = document.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_run = logo_para.add_run()
        logo_run.add_picture("jiit_logo.png", width=Inches(2.0))
    except Exception as e:
        error_para = document.add_paragraph(f"[Logo 'jiit_logo.png' not found: {e}]")
        error_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add spacing
    document.add_paragraph()
    
    # Project Synopsis title
    synopsis_para = document.add_paragraph("Project Synopsis")
    synopsis_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    synopsis_run = synopsis_para.runs[0] if synopsis_para.runs else synopsis_para.add_run()
    synopsis_run.bold = True
    synopsis_run.font.size = Pt(16)
    
    # Project Title
    project_title_para = document.add_paragraph(f"TITLE: {user_data['title']}")
    project_title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    project_title_run = project_title_para.runs[0] if project_title_para.runs else project_title_para.add_run()
    project_title_run.bold = True
    project_title_run.font.size = Pt(14)
    
    # Add spacing
    document.add_paragraph()
    
    # SUBMITTED BY
    submitted_para = document.add_paragraph("SUBMITTED BY:")
    submitted_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    submitted_run = submitted_para.runs[0] if submitted_para.runs else submitted_para.add_run()
    submitted_run.bold = True
    submitted_run.font.size = Pt(14)
    
    # Team Members Table
    table = document.add_table(rows=1, cols=2)
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = "NAME"
    header_cells[0].paragraphs[0].runs[0].bold = True
    header_cells[1].text = "ENROLLMENT"
    header_cells[1].paragraphs[0].runs[0].bold = True
    
    # Add team members
    for member in user_data['team']:
        row_cells = table.add_row().cells
        row_cells[0].text = member['name']
        row_cells[1].text = member['enroll']
    
    # Add spacing
    document.add_paragraph()
    
    # Submitted to
    submitted_to_para = document.add_paragraph("Submitted to:")
    submitted_to_run = submitted_to_para.runs[0] if submitted_to_para.runs else submitted_to_para.add_run()
    submitted_to_run.bold = True
    submitted_to_run.font.size = Pt(14)
    
    # Add guides
    for guide in user_data['guides']:
        guide_para = document.add_paragraph(guide.strip())
        guide_para.style = 'List Bullet'
    
    # Page break
    document.add_page_break()
    
    # === AI Content Pages ===
    
    # OBJECTIVE
    document.add_heading("OBJECTIVE:", level=2)
    objective_para = document.add_paragraph(ai_data['objective'])
    objective_para.paragraph_format.space_after = Pt(6)
    
    # INTRODUCTION
    document.add_heading("INTRODUCTION:", level=2)
    intro_para = document.add_paragraph(ai_data['introduction'])
    intro_para.paragraph_format.space_after = Pt(6)
    
    # KEY FEATURES
    document.add_heading("KEY FEATURES:", level=2)
    for feature in ai_data['key_features']:
        feature_para = document.add_paragraph(feature, style='List Bullet')
        # Reduce spacing and indentation for list items
        feature_para.paragraph_format.space_before = Pt(0)
        feature_para.paragraph_format.space_after = Pt(2)
        feature_para.paragraph_format.line_spacing = Pt(11)
        feature_para.paragraph_format.left_indent = Inches(0.2)  # Reduced left indent
        feature_para.paragraph_format.first_line_indent = Inches(-0.15)  # Negative to bring text closer to bullet
    
    # TECHNOLOGIES USED
    document.add_heading("TECHNOLOGIES USED:", level=2)
    for category, items in ai_data['technologies_used'].items():
        # Add category as bold paragraph
        category_para = document.add_paragraph()
        category_run = category_para.add_run(category)
        category_run.bold = True
        category_para.paragraph_format.space_after = Pt(3)
        
        # Add items as bulleted list
        if isinstance(items, list):
            for item in items:
                item_para = document.add_paragraph(str(item), style='List Bullet 2')
                item_para.paragraph_format.space_before = Pt(0)
                item_para.paragraph_format.space_after = Pt(2)
                item_para.paragraph_format.line_spacing = Pt(11)
                item_para.paragraph_format.left_indent = Inches(0.3)  # Reduced left indent
                item_para.paragraph_format.first_line_indent = Inches(-0.15)  # Negative to bring text closer to bullet
        else:
            item_para = document.add_paragraph(str(items), style='List Bullet 2')
            item_para.paragraph_format.space_before = Pt(0)
            item_para.paragraph_format.space_after = Pt(2)
            item_para.paragraph_format.line_spacing = Pt(11)
            item_para.paragraph_format.left_indent = Inches(0.3)
            item_para.paragraph_format.first_line_indent = Inches(-0.15)
    
    # REFERENCES
    document.add_heading("REFERENCES:", level=2)
    for category, items in ai_data['references'].items():
        # Add category as bold paragraph
        category_para = document.add_paragraph()
        category_run = category_para.add_run(category)
        category_run.bold = True
        category_para.paragraph_format.space_after = Pt(3)
        
        # Add items as bulleted list
        if isinstance(items, list):
            for item in items:
                item_para = document.add_paragraph(str(item), style='List Bullet 2')
                item_para.paragraph_format.space_before = Pt(0)
                item_para.paragraph_format.space_after = Pt(2)
                item_para.paragraph_format.line_spacing = Pt(11)
                item_para.paragraph_format.left_indent = Inches(0.3)  # Reduced left indent
                item_para.paragraph_format.first_line_indent = Inches(-0.15)  # Negative to bring text closer to bullet
        else:
            item_para = document.add_paragraph(str(items), style='List Bullet 2')
            item_para.paragraph_format.space_before = Pt(0)
            item_para.paragraph_format.space_after = Pt(2)
            item_para.paragraph_format.line_spacing = Pt(11)
            item_para.paragraph_format.left_indent = Inches(0.3)
            item_para.paragraph_format.first_line_indent = Inches(-0.15)
    
    # === Image Pages (Each on separate page) ===
    if uploaded_images:
        for img_file in uploaded_images:
            # Add page break
            document.add_page_break()
            
            # Add heading using filename (centered)
            heading = img_file.name.split('.')[0].replace('_', ' ').title()
            heading_para = document.add_heading(heading, level=2)
            heading_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add spacing
            document.add_paragraph()
            
            # Add image, centered and scaled
            try:
                img_file.seek(0)
                # Read the image bytes from UploadedFile
                img_bytes = img_file.read()
                img_file.seek(0)  # Reset for potential reuse
                
                # Create a BytesIO object from the image bytes
                img_io = io.BytesIO(img_bytes)
                
                img_para = document.add_paragraph()
                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_run = img_para.add_run()
                img_run.add_picture(img_io, width=Inches(6.0))
            except Exception as e:
                error_para = document.add_paragraph(f"[Error loading image: {e}]")
                error_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save to BytesIO buffer
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer

# --- (MAIN STREAMLIT UI FUNCTION) ---

def show():
    st.title("🤖 JIIT Project Synopsis Generator")
    
    with st.form("synopsis_form"):
        st.header("1. Cover Page Details")
        
        project_title = st.text_input("Project Title*", placeholder="e.g., Student Management System")
        
        project_category = st.text_input("Project Category/Lab*", placeholder="e.g., Open Source Lab Project")
        
        # --- NEW PRECISE INPUT FOR TEAM ---
        st.subheader("Team Members")
        member_count = st.number_input("How many members in your team?", min_value=1, max_value=6, value=1)
        
        team_data = []
        for i in range(member_count):
            cols = st.columns(2)
            with cols[0]:
                name = st.text_input(f"Member {i+1} Name", key=f"name_{i}")
            with cols[1]:
                enroll = st.text_input(f"Member {i+1} Enrollment", key=f"enroll_{i}")
            team_data.append({"name": name, "enroll": enroll})
        # --- END NEW SECTION ---
            
        st.subheader("Guides")
        guides = st.text_area("Submitted to (Guides)*", 
                              placeholder="Dr. Alkha Seghal\nDr. Tribhuvan Tiwary")

        st.header("2. Additional Project Images (Optional)")
        uploaded_images = st.file_uploader(
            "Upload Diagrams, Screenshots, etc. (One page per image)",
            type=["png", "jpg", "jpeg"],
            accept_multiple_files=True 
        )
        
        st.header("3. Generate Document")
        submitted = st.form_submit_button("✨ Generate Project PDF")
        
    if submitted:
        # 1. Validate input
        if not all([project_title,project_category, guides]) or not all(m['name'] and m['enroll'] for m in team_data):
            st.error("Please fill in all required fields (*), including all team member names and enrollments.")
        else:
            # 2. Store user data
            user_data = {
                "title": project_title,
                "category": project_category,
                "team": team_data, # Use the new list of dictionaries
                "guides": guides.split('\n')
            }
            
            # 3. Call REAL AI function
            with st.spinner("🤖 Contacting JiitAi... Generating unique project content..."):
                ai_data = call_gemini_ai(project_title)
            
            if ai_data:
                st.success("AI Content Generated!")
                
                # 4. Generate both PDF and DOCX
                with st.spinner("📄 Building your PDF..."):
                    pdf_buffer = create_project_pdf(user_data, ai_data, uploaded_images)
                
                with st.spinner("📝 Building your DOCX..."):
                    doc_buffer = create_project_doc(user_data, ai_data, uploaded_images)
                
                # 5. Provide downloads in two columns
                col1, col2 = st.columns(2)
                
                with col1:
                    st.download_button(
                        label="⬇️ Download Synopsis PDF",
                        data=pdf_buffer,
                        file_name=f"{project_title.replace(' ', '_')}_Synopsis.pdf",
                        mime="application/pdf",
                        use_container_width=True
                    )
                
                with col2:
                    st.download_button(
                        label="⬇️ Download Synopsis .docx",
                        data=doc_buffer,
                        file_name=f"{project_title.replace(' ', '_')}_Synopsis.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
            else:
                st.error("Failed to generate documents. Could not get AI content.")

# --- (APP EXECUTION) ---

if __name__ == "__main__":
    show()